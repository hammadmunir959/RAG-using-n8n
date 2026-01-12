"""
Document text extraction and chunking utilities.
Supports PDF, DOCX, TXT, CSV, and JSON files.
"""
import io
import json
import csv
import logging
from typing import List, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text(filename: str, content: bytes) -> str:
    """
    Extract text content from a document based on file type.
    
    Args:
        filename: Original filename with extension
        content: Raw file bytes
        
    Returns:
        Extracted text content as string
    """
    ext = Path(filename).suffix.lower()
    
    try:
        if ext == '.pdf':
            return extract_pdf(content)
        elif ext == '.docx':
            return extract_docx(content)
        elif ext == '.txt':
            return extract_txt(content)
        elif ext == '.csv':
            return extract_csv(content)
        elif ext == '.json':
            return extract_json(content)
        else:
            logger.warning(f"Unsupported file type: {ext}, attempting plain text")
            return extract_txt(content)
    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {e}")
        raise


def extract_pdf(content: bytes) -> str:
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            except Exception as e:
                logger.warning(f"Error extracting page {page_num + 1}: {e}")
                continue
        
        return "\n\n".join(text_parts)
    except ImportError:
        logger.error("pypdf not installed. Run: pip install pypdf")
        raise


def extract_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        raise


def extract_txt(content: bytes) -> str:
    """Extract text from plain text file with encoding detection."""
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'ascii']
    
    for encoding in encodings:
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    
    # Fallback: decode with errors ignored
    return content.decode('utf-8', errors='ignore')


def extract_csv(content: bytes) -> str:
    """Extract text from CSV, converting to readable format."""
    text = extract_txt(content)
    lines = text.strip().split('\n')
    
    if not lines:
        return ""
    
    try:
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        
        if not rows:
            return text
        
        # Use first row as headers
        headers = rows[0] if rows else []
        result_parts = []
        
        for i, row in enumerate(rows[1:], 1):
            row_data = []
            for j, value in enumerate(row):
                if value.strip():
                    header = headers[j] if j < len(headers) else f"Column {j+1}"
                    row_data.append(f"{header}: {value}")
            if row_data:
                result_parts.append(f"[Row {i}] " + ", ".join(row_data))
        
        return "\n".join(result_parts)
    except csv.Error:
        return text


def extract_json(content: bytes) -> str:
    """Extract text from JSON, flattening structure."""
    text = extract_txt(content)
    
    try:
        data = json.loads(text)
        return flatten_json(data)
    except json.JSONDecodeError:
        return text


def flatten_json(obj, prefix: str = "") -> str:
    """Recursively flatten JSON to readable text."""
    parts = []
    
    if isinstance(obj, dict):
        for key, value in obj.items():
            new_prefix = f"{prefix}.{key}" if prefix else key
            parts.append(flatten_json(value, new_prefix))
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            new_prefix = f"{prefix}[{i}]"
            parts.append(flatten_json(item, new_prefix))
    else:
        if obj is not None and str(obj).strip():
            return f"{prefix}: {obj}"
        return ""
    
    return "\n".join(p for p in parts if p)


def chunk_text(
    text: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    separator: str = "\n\n"
) -> List[str]:
    """
    Split text into overlapping chunks for embedding.
    
    Args:
        text: Full document text
        chunk_size: Target size of each chunk in characters
        chunk_overlap: Overlap between consecutive chunks
        separator: Primary separator to split on
        
    Returns:
        List of text chunks
    """
    if not text or not text.strip():
        return []
    
    # Clean the text
    text = text.strip()
    
    # If text is small enough, return as single chunk
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    
    # First, try to split by paragraphs
    paragraphs = text.split(separator)
    
    current_chunk = ""
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        # If adding this paragraph would exceed chunk size
        if len(current_chunk) + len(para) + len(separator) > chunk_size:
            if current_chunk:
                chunks.append(current_chunk.strip())
                # Start new chunk with overlap from previous
                if chunk_overlap > 0 and len(current_chunk) > chunk_overlap:
                    current_chunk = current_chunk[-chunk_overlap:] + separator + para
                else:
                    current_chunk = para
            else:
                # Single paragraph exceeds chunk size, split by sentences
                sentence_chunks = split_long_text(para, chunk_size, chunk_overlap)
                chunks.extend(sentence_chunks[:-1])
                current_chunk = sentence_chunks[-1] if sentence_chunks else ""
        else:
            if current_chunk:
                current_chunk += separator + para
            else:
                current_chunk = para
    
    # Add the last chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks


def split_long_text(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    """Split long text that doesn't have paragraph breaks."""
    chunks = []
    
    # Try to split by sentences
    sentences = text.replace('. ', '.|').replace('? ', '?|').replace('! ', '!|').split('|')
    
    current = ""
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        
        if len(current) + len(sentence) + 1 > chunk_size:
            if current:
                chunks.append(current)
                if chunk_overlap > 0:
                    current = current[-chunk_overlap:] + " " + sentence
                else:
                    current = sentence
            else:
                # Single sentence too long, hard split
                chunks.append(sentence[:chunk_size])
                current = sentence[chunk_size - chunk_overlap:] if chunk_overlap else ""
        else:
            current = current + " " + sentence if current else sentence
    
    if current.strip():
        chunks.append(current.strip())
    
    return chunks


def get_document_metadata(filename: str, content: bytes) -> dict:
    """Extract metadata from document."""
    ext = Path(filename).suffix.lower()
    
    metadata = {
        "filename": filename,
        "file_type": ext.lstrip('.'),
        "file_size": len(content),
        "char_count": 0,
        "estimated_pages": 0
    }
    
    try:
        text = extract_text(filename, content)
        metadata["char_count"] = len(text)
        # Rough page estimate (assuming ~3000 chars per page)
        metadata["estimated_pages"] = max(1, len(text) // 3000)
    except:
        pass
    
    return metadata
